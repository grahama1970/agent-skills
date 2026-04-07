"""Paper parsing: Markdown, LaTeX, plain text, and DOCX to PaperModel."""
from __future__ import annotations

import re
import tempfile
from pathlib import Path

from docx import Document

from .models import PaperModel, ParagraphModel, SectionModel, SentenceModel

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_SENTENCE_SPLIT_RE = re.compile(r"(?<=[.!?])\s+")

# LaTeX heading patterns
_LATEX_SECTION_RE = re.compile(r"\\(section|subsection|subsubsection)\*?\{(.+?)\}")
_LATEX_TITLE_RE = re.compile(r"\\title\{(.+?)\}")
_LATEX_ABSTRACT_BEGIN = re.compile(r"\\begin\{abstract\}")
_LATEX_ABSTRACT_END = re.compile(r"\\end\{abstract\}")

# LaTeX markup to strip for text analysis (order matters)
_LATEX_STRIP_PATTERNS = [
    # Remove block environments (equations, figures, tables, etc.)
    (re.compile(r"\\begin\{(equation|align|figure|table|tabular|itemize|enumerate|lstlisting|verbatim|thebibliography)\*?\}.*?\\end\{\1\*?\}", re.DOTALL), ""),
    # Unwrap text formatting commands (keep inner text)
    (re.compile(r"\\(textbf|textit|emph|texttt|textsf|textrm|underline)\{([^}]*)\}"), r"\2"),
    # Remove reference commands
    (re.compile(r"\\(cite|ref|label|eqref|pageref|autoref|cref)\{[^}]*\}"), ""),
    # Unwrap footnotes (keep text)
    (re.compile(r"\\(footnote|footnotetext)\{([^}]*)\}"), r"\2"),
    # Remove inline math
    (re.compile(r"\$[^$]+\$"), ""),
    # Remove remaining commands
    (re.compile(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{[^}]*\})?"), ""),
    # Remove stray braces
    (re.compile(r"[{}]"), ""),
    # Remove LaTeX comments
    (re.compile(r"%.*$", re.MULTILINE), ""),
    # Collapse whitespace
    (re.compile(r"\s+"), " "),
]


def _strip_latex(text: str) -> str:
    """Remove LaTeX markup, returning plain text for analysis."""
    for pattern, replacement in _LATEX_STRIP_PATTERNS:
        text = pattern.sub(replacement, text)
    return text.strip()


def _sentence_models(paragraph_id: str, text: str) -> list[SentenceModel]:
    chunks = [s.strip() for s in _SENTENCE_SPLIT_RE.split(text.strip()) if s.strip()]
    if not chunks and text.strip():
        chunks = [text.strip()]
    return [
        SentenceModel(
            id=f"{paragraph_id}:s{i+1}",
            text=chunk,
            word_count=len(chunk.split()),
        )
        for i, chunk in enumerate(chunks)
    ]


def _paragraph_model(section_id: str, index: int, text: str) -> ParagraphModel:
    text = text.strip()
    pid = f"{section_id}:p{index}"
    return ParagraphModel(
        id=pid,
        text=text,
        sentences=_sentence_models(pid, text),
        word_count=len(text.split()),
    )


def _finalize_section(
    sections: list[SectionModel], title: str, level: int, paragraphs: list[str],
) -> None:
    sid = f"sec_{len(sections)+1}"
    paragraph_models = [
        _paragraph_model(sid, i + 1, p) for i, p in enumerate(paragraphs) if p.strip()
    ]
    if paragraph_models:
        sections.append(
            SectionModel(id=sid, title=title, level=level, paragraphs=paragraph_models),
        )


def parse_latex(path: Path) -> PaperModel:
    """Parse a LaTeX file into sections, stripping markup for analysis."""
    raw_text = path.read_text(encoding="utf-8")

    # Extract title
    title_match = _LATEX_TITLE_RE.search(raw_text)
    title = _strip_latex(title_match.group(1)) if title_match else path.stem.replace("_", " ").title()

    # Extract abstract
    abstract: str | None = None
    abs_begin = _LATEX_ABSTRACT_BEGIN.search(raw_text)
    abs_end = _LATEX_ABSTRACT_END.search(raw_text)
    if abs_begin and abs_end:
        abstract = _strip_latex(raw_text[abs_begin.end():abs_end.start()])

    # Parse sections
    lines = raw_text.splitlines()
    sections: list[SectionModel] = []
    current_title = "Preamble"
    current_level = 0
    paragraph_accum: list[str] = []
    buffer: list[str] = []
    in_abstract = False
    in_preamble = True

    level_map = {"section": 1, "subsection": 2, "subsubsection": 3}

    def flush_paragraph():
        nonlocal paragraph_accum, buffer
        joined = " ".join(paragraph_accum).strip()
        if joined:
            buffer.append(joined)
        paragraph_accum = []

    for line in lines:
        stripped = line.strip()

        # Skip preamble until \begin{document}
        if in_preamble:
            if stripped == "\\begin{document}":
                in_preamble = False
            continue
        if stripped == "\\end{document}":
            break
        if stripped.startswith("\\maketitle"):
            continue

        # Track abstract
        if _LATEX_ABSTRACT_BEGIN.search(stripped):
            in_abstract = True
            flush_paragraph()
            if buffer:
                _finalize_section(sections, current_title, current_level, buffer)
                buffer = []
            current_title = "Abstract"
            current_level = 1
            continue
        if _LATEX_ABSTRACT_END.search(stripped):
            in_abstract = False
            flush_paragraph()
            if buffer:
                _finalize_section(sections, current_title, current_level, buffer)
                buffer = []
            current_title = "Body"
            current_level = 1
            continue

        # Detect section headings
        sec_match = _LATEX_SECTION_RE.search(stripped)
        if sec_match:
            flush_paragraph()
            if buffer:
                _finalize_section(sections, current_title, current_level, buffer)
                buffer = []
            current_level = level_map.get(sec_match.group(1), 1)
            current_title = _strip_latex(sec_match.group(2))
            continue

        # Skip comments
        if stripped.startswith("%"):
            continue
        # Skip environment markers, bibliography
        if stripped.startswith("\\begin{") or stripped.startswith("\\end{"):
            continue
        if stripped.startswith("\\bibliography") or stripped.startswith("\\bibliographystyle"):
            continue

        # Empty line = paragraph break
        if not stripped:
            flush_paragraph()
            continue

        # Strip LaTeX from content and accumulate
        clean = _strip_latex(stripped)
        if clean:
            paragraph_accum.append(clean)

    flush_paragraph()
    if buffer:
        _finalize_section(sections, current_title, current_level, buffer)

    # Build clean raw_text for metrics
    clean_raw = _strip_latex(raw_text)

    return PaperModel(
        source_path=str(path),
        title=title,
        abstract=abstract,
        sections=sections or [
            SectionModel(id="sec_1", title="Body", paragraphs=[], level=1),
        ],
        raw_text=clean_raw,
    )


def parse_markdown_or_text(path: Path) -> PaperModel:
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    title = path.stem.replace("_", " ").title()
    sections: list[SectionModel] = []
    current_title = "Introduction"
    current_level = 1
    buffer: list[str] = []
    abstract: str | None = None

    if lines and lines[0].startswith("# "):
        title = lines[0][2:].strip()
        lines = lines[1:]

    paragraph_accum: list[str] = []

    def flush_paragraph_buffer() -> None:
        nonlocal paragraph_accum, buffer
        joined = " ".join([x.strip() for x in paragraph_accum if x.strip()]).strip()
        if joined:
            buffer.append(joined)
        paragraph_accum = []

    for line in lines:
        heading = _HEADING_RE.match(line)
        if heading:
            flush_paragraph_buffer()
            if buffer:
                _finalize_section(sections, current_title, current_level, buffer)
            current_level = len(heading.group(1))
            current_title = heading.group(2).strip()
            buffer = []
            continue

        if not line.strip():
            flush_paragraph_buffer()
            continue

        paragraph_accum.append(line.strip())

    flush_paragraph_buffer()
    if buffer:
        _finalize_section(sections, current_title, current_level, buffer)

    if sections and sections[0].title.lower() == "abstract" and sections[0].paragraphs:
        abstract = "\n\n".join(p.text for p in sections[0].paragraphs)

    return PaperModel(
        source_path=str(path),
        title=title,
        abstract=abstract,
        sections=sections or [
            SectionModel(id="sec_1", title="Body", paragraphs=[], level=1),
        ],
        raw_text=text,
    )


def parse_docx(path: Path) -> PaperModel:
    doc = Document(str(path))
    lines = [p.text.rstrip() for p in doc.paragraphs]
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".txt", encoding="utf-8", delete=False,
    ) as tmp:
        tmp.write("\n".join(lines))
        tmp_path = Path(tmp.name)
    try:
        return parse_markdown_or_text(tmp_path)
    finally:
        tmp_path.unlink(missing_ok=True)


def parse_paper(path: Path) -> PaperModel:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return parse_docx(path)
    if suffix in (".tex", ".latex"):
        return parse_latex(path)
    return parse_markdown_or_text(path)
