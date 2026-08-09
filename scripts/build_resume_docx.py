"""Build an ATS-oriented DOCX from the same Markdown the PDF is built from.

Measured 2026 parse rates put DOCX ahead of a text PDF across the major
trackers — roughly 97% against 91% on average, and 97% against 83% on Taleo,
which is the widest gap and the oldest parser. The PDF stays the artifact a
human reads; this is the one to upload when a form offers the choice.

Deliberately plain: single column, conventional headings, no tables, no text
boxes, no images, standard date lines. Everything a parser trips over is
absent, and the typography that makes the PDF worth looking at is not
reproduced here because it earns nothing in a parser and costs extraction
accuracy.

Built from RESUME.md through the same ``--omit-section`` convention as the PDF,
so the three exports cannot drift apart.

    uv run --with python-docx python scripts/build_resume_docx.py \\
      RESUME.md docs/resume/graham-anderson-resume.docx \\
      --omit-section "DEEPER DETAIL"
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

# Reuse the PDF builder's section filter so one convention governs both exports.
sys.path.insert(0, str(Path(__file__).resolve().parent))
from build_markdown_pdf import drop_sections  # noqa: E402

LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")


def plain(text: str) -> str:
    """Markdown inline markup reduced to the words a parser will index.

    Link labels are already the visible URL in this resume, so the target adds
    nothing a parser can use and its removal keeps lines short.
    """
    text = LINK_RE.sub(r"\1", text)
    text = BOLD_RE.sub(r"\1", text)
    return text.replace("`", "").strip()


def build(md: str, out: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    # One clean base style; parsers key off heading levels, not fonts.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    for raw in md.splitlines():
        line = raw.strip()
        if not line or line == "<!-- pdf-only -->":
            continue
        if line.startswith("# ") and not line.startswith("##"):
            doc.add_heading(plain(line[2:]), level=0)
        elif line.startswith("## "):
            doc.add_heading(plain(line[3:]), level=1)
        elif line.startswith("### "):
            doc.add_heading(plain(line[4:]), level=2)
        elif line.startswith("> "):
            doc.add_paragraph(plain(line[2:]))
        elif line.startswith("- "):
            doc.add_paragraph(plain(line[2:]), style="List Bullet")
        else:
            doc.add_paragraph(plain(line))

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))


def main(argv: list[str] | None = None) -> int:
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("markdown", type=Path)
    p.add_argument("output", type=Path)
    p.add_argument("--omit-section", action="append", default=[], dest="omit_sections")
    args = p.parse_args(sys.argv[1:] if argv is None else argv)

    try:
        if not args.markdown.is_file():
            raise FileNotFoundError(f"Markdown file does not exist: {args.markdown}")
        text = args.markdown.read_text(encoding="utf-8")
        if not text.strip():
            raise ValueError(f"Markdown file is empty: {args.markdown}")
        build(drop_sections(text, tuple(args.omit_sections)), args.output)
        size = args.output.stat().st_size
        if size < 2000:
            raise ValueError(f"Generated DOCX is unexpectedly small: {size} bytes")
    except Exception as exc:  # noqa: BLE001 — CLI boundary
        print(f"error: {exc}", file=sys.stderr)
        return 1
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
